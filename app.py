"""
Word 格式化工具 MVP v2
用户自主选择标题，不自动识别
"""

import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

# 页面配置
st.set_page_config(
    page_title="Word 格式化工具",
    page_icon="📝",
    layout="wide"
)

# 字体选项
FONT_OPTIONS = {
    "宋体": "宋体",
    "黑体": "黑体", 
    "楷体": "楷体",
    "仿宋": "仿宋",
    "Times New Roman": "Times New Roman",
    "Arial": "Arial"
}

# 字号选项（磅值）
SIZE_OPTIONS = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5, "七号": 5.5, "八号": 5
}

def clear_paragraph_format(para):
    """清除段落的格式"""
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None
    para.paragraph_format.right_indent = None
    
    for run in para.runs:
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False

def set_run_font(run, font_name, font_size, bold=False):
    """设置run的字体（包括中文字体）"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = False
    run.font.underline = False
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def apply_format(doc, format_rules, title_selections):
    """应用格式规则到文档"""
    
    # 首先清除所有段落格式
    for para in doc.paragraphs:
        clear_paragraph_format(para)
    
    # 记录哪些段落是标题
    title_indices = set()
    
    # 处理标题
    for selection in title_selections:
        idx = selection['index']
        level = selection['level']
        title_indices.add(idx)
        
        para = doc.paragraphs[idx]
        
        # 根据级别获取格式
        if level == 1:
            font_name = format_rules['title1_font']
            font_size = format_rules['title1_size']
            bold = format_rules['title1_bold']
            align = format_rules['title1_align']
        elif level == 2:
            font_name = format_rules['title2_font']
            font_size = format_rules['title2_size']
            bold = format_rules['title2_bold']
            align = format_rules['title2_align']
        else:
            font_name = format_rules['title3_font']
            font_size = format_rules['title3_size']
            bold = format_rules['title3_bold']
            align = format_rules['title3_align']
        
        # 应用格式
        for run in para.runs:
            set_run_font(run, font_name, font_size, bold)
        
        # 对齐方式
        if align == 'center':
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 应用正文格式到非标题段落
    body_font = format_rules['body_font']
    body_size = format_rules['body_size']
    body_bold = format_rules['body_bold']
    
    for i, para in enumerate(doc.paragraphs):
        if i in title_indices:
            continue
        
        if not para.text.strip():
            continue
        
        # 应用正文格式
        for run in para.runs:
            set_run_font(run, body_font, body_size, body_bold)
        
        # 段落格式
        if format_rules.get('first_indent'):
            para.paragraph_format.first_line_indent = Cm(0.74)
        
        if format_rules.get('line_spacing'):
            para.paragraph_format.line_spacing = format_rules['line_spacing']
        
        if format_rules.get('align_justify'):
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    return doc

# 主界面
st.markdown('<h1 style="font-size:2rem;color:#1f77b4;">📝 Word 格式化工具</h1>', unsafe_allow_html=True)
st.markdown("自定义格式规则，一键排版 Word 文档")

# 初始化session state
if 'doc' not in st.session_state:
    st.session_state.doc = None
if 'paragraphs' not in st.session_state:
    st.session_state.paragraphs = []
if 'selected_titles' not in st.session_state:
    st.session_state.selected_titles = {}
if 'uploaded_file_name' not in st.session_state:
    st.session_state.uploaded_file_name = None

# 步骤1：上传文档
st.markdown("---")
st.markdown("### 📄 步骤1：上传文档")

uploaded_file = st.file_uploader("选择 Word 文档（.docx 格式）", type="docx", key="file_uploader")

if uploaded_file and uploaded_file.name != st.session_state.uploaded_file_name:
    st.success(f"已上传: {uploaded_file.name}")
    
    # 读取文档
    doc = Document(uploaded_file)
    st.info(f"文档共有 {len(doc.paragraphs)} 个段落")
    
    # 提取段落
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                'index': i,
                'text': text,
                'short_text': text[:60] + ('...' if len(text) > 60 else '')
            })
    
    st.session_state.doc = doc
    st.session_state.paragraphs = paragraphs
    st.session_state.selected_titles = {}
    st.session_state.uploaded_file_name = uploaded_file.name

# 步骤2：设置格式
if st.session_state.doc:
    st.markdown("---")
    st.markdown("### ⚙️ 步骤2：设置格式规则")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("**一级标题**")
        title1_font = st.selectbox("字体", list(FONT_OPTIONS.keys()), key="t1_font", index=1)
        title1_size = st.selectbox("字号", list(SIZE_OPTIONS.keys()), key="t1_size", index=4)
        title1_bold = st.checkbox("加粗", value=True, key="t1_bold")
        title1_align = st.radio("对齐", ["居中", "左对齐"], key="t1_align", horizontal=True)
    
    with col2:
        st.markdown("**二级标题**")
        title2_font = st.selectbox("字体", list(FONT_OPTIONS.keys()), key="t2_font", index=2)
        title2_size = st.selectbox("字号", list(SIZE_OPTIONS.keys()), key="t2_size", index=6)
        title2_bold = st.checkbox("加粗", value=True, key="t2_bold")
        title2_align = st.radio("对齐", ["左对齐", "居中"], key="t2_align", horizontal=True)
    
    with col3:
        st.markdown("**三级标题**")
        title3_font = st.selectbox("字体", list(FONT_OPTIONS.keys()), key="t3_font", index=3)
        title3_size = st.selectbox("字号", list(SIZE_OPTIONS.keys()), key="t3_size", index=8)
        title3_bold = st.checkbox("加粗", value=True, key="t3_bold")
        title3_align = st.radio("对齐", ["左对齐", "居中"], key="t3_align", horizontal=True)
    
    st.markdown("---")
    st.markdown("**正文格式**")
    col_body1, col_body2, col_body3 = st.columns(3)
    
    with col_body1:
        body_font = st.selectbox("字体", list(FONT_OPTIONS.keys()), key="body_font", index=0)
        body_size = st.selectbox("字号", list(SIZE_OPTIONS.keys()), key="body_size", index=6)
    
    with col_body2:
        line_spacing = st.slider("行距", 1.0, 3.0, 1.0, 0.5, key="line_spacing")
        first_indent = st.checkbox("首行缩进2字符", value=True, key="first_indent")
    
    with col_body3:
        align_justify = st.checkbox("两端对齐", value=True, key="align_justify")
        body_bold = st.checkbox("加粗", value=False, key="body_bold")

    # 步骤3：选择标题
    st.markdown("---")
    st.markdown("### ✅ 步骤3：选择标题段落")
    st.markdown("勾选需要设为标题的段落，并选择标题级别：")
    
    # 显示所有段落供选择
    for p in st.session_state.paragraphs:
        idx = p['index']
        col_check, col_text, col_level = st.columns([0.5, 5, 1.5])
        
        with col_check:
            is_title = st.checkbox("", value=idx in st.session_state.selected_titles, key=f"check_{idx}")
        
        with col_text:
            st.markdown(f"`{p['short_text']}`")
        
        with col_level:
            if is_title:
                level = st.selectbox("级别", [1, 2, 3], 
                    index=st.session_state.selected_titles.get(idx, {}).get('level', 0) - 1 if idx in st.session_state.selected_titles else 0,
                    key=f"level_{idx}", label_visibility="collapsed")
                st.session_state.selected_titles[idx] = {'level': level, 'text': p['text']}
            else:
                st.markdown("_正文_")
                if idx in st.session_state.selected_titles:
                    del st.session_state.selected_titles[idx]
    
    # 显示选择结果
    if st.session_state.selected_titles:
        st.info(f"已选择 {len(st.session_state.selected_titles)} 个标题段落")
    else:
        st.warning("未选择任何标题，所有段落将应用正文格式")

    # 步骤4：生成文档
    st.markdown("---")
    st.markdown("### 📥 步骤4：生成并下载")
    
    if st.button("🚀 开始排版", type="primary"):
        # 整理格式规则
        format_rules = {
            'title1_font': FONT_OPTIONS[title1_font],
            'title1_size': SIZE_OPTIONS[title1_size],
            'title1_bold': title1_bold,
            'title1_align': 'center' if title1_align == "居中" else 'left',
            'title2_font': FONT_OPTIONS[title2_font],
            'title2_size': SIZE_OPTIONS[title2_size],
            'title2_bold': title2_bold,
            'title2_align': 'left' if title2_align == "左对齐" else 'center',
            'title3_font': FONT_OPTIONS[title3_font],
            'title3_size': SIZE_OPTIONS[title3_size],
            'title3_bold': title3_bold,
            'title3_align': 'left' if title3_align == "左对齐" else 'center',
            'body_font': FONT_OPTIONS[body_font],
            'body_size': SIZE_OPTIONS[body_size],
            'body_bold': body_bold,
            'line_spacing': line_spacing,
            'first_indent': first_indent,
            'align_justify': align_justify
        }
        
        # 准备标题选择数据
        title_selections = [
            {'index': idx, 'level': info['level']}
            for idx, info in st.session_state.selected_titles.items()
        ]
        
        # 复制文档并应用格式
        output_doc = Document(uploaded_file)
        output_doc = apply_format(output_doc, format_rules, title_selections)
        
        # 保存到内存
        output = BytesIO()
        output_doc.save(output)
        output.seek(0)
        
        st.success("排版完成！")
        
        st.download_button(
            label="📥 下载排版后的文档",
            data=output,
            file_name=f"排版完成_{uploaded_file.name}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# 页脚
st.markdown("---")
st.markdown("💡 **提示**：排版后会清除原有的加粗、斜体、下划线等格式，统一应用新格式")
